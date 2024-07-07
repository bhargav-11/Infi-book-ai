import base64
from io import BytesIO

import os
from doc2docx import convert
from docx import Document
import docx
import markdown2
from htmldocx import HtmlToDocx
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


def generate_document(text, idx,title):
  # Convert the Markdown content to HTML using markdown2
  html = markdown2.markdown(text, extras=["tables"])

  # Create a new Word document
  doc = docx.Document()

  # Add the title to the document
  doc.add_heading(title, 0)

  # Convert HTML to Word document using HtmlToDocx
  html_to_docx_parser = HtmlToDocx()
  html_to_docx_parser.add_html_to_document(html, doc)

  # Save the document to a BytesIO object
  file_bytes = BytesIO()
  doc.save(file_bytes)
  file_bytes.seek(0)
  
  text_bytes = file_bytes.read()

  return text_bytes

def generate_download_link(text,title):
  # Convert the Markdown content to HTML using markdown2
  html = markdown2.markdown(text, extras=["tables"])

  # Create a new Word document
  doc = docx.Document()

  # Add the title to the document
  doc.add_heading(title, 0)

  # Convert HTML to Word document using HtmlToDocx
  html_to_docx_parser = HtmlToDocx()
  html_to_docx_parser.add_html_to_document(html, doc)


  

  # Add borders to all tables in the document
  for table in doc.tables:
        print("Table :",table)
        set_table_border(table, border_size=4)
        
        # Optionally, set a background color for header row
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '2F5496')
            cell._tc.get_or_add_tcPr().append(shading_elm)
    
  # Save the document to a BytesIO object
  file_bytes = BytesIO()
  doc.save(file_bytes)
  file_bytes.seek(0)

  # Create a download link for the Word document
  download_link = create_download_link(file_bytes.read(),
                                       f"{title}.docx")
  return download_link


def create_download_link(val, filename):
  b64 = base64.b64encode(val)
  return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64.decode()}" download="{filename}">Download file</a>'


def extract_doc_text(uploaded_file):
    
  tmpdirname = './tmp'
  if not os.path.exists(tmpdirname):
      os.makedirs(tmpdirname)

  # Define the paths for the input and output files
  input_file_path = os.path.join(tmpdirname, uploaded_file.name)
  output_file_path = os.path.join(tmpdirname, "output.docx")
  
  # Save the uploaded file to the tmp location
  with open(input_file_path, 'wb') as f:
      f.write(uploaded_file.read())
  
  try:
      convert(input_file_path, output_file_path)
      all_text = extract_docx_text(output_file_path)      
      return all_text  
  finally:
      # Clean up the temporary files if needed
      if os.path.exists(input_file_path):
          os.remove(input_file_path)
      if os.path.exists(output_file_path):
          os.remove(output_file_path)

def extract_docx_text(docx_file):
    doc = Document(docx_file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)


def set_table_border(table, border_size=1):
    tbl_pr = table._element.xpath('w:tblPr')[0]
    borders = OxmlElement('w:tblBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        edge = OxmlElement(f'w:{border}')
        edge.set(qn('w:sz'), str(border_size))
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:color'), '000000')
        borders.append(edge)
    tbl_pr.append(borders)